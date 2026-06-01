#!/usr/bin/env python3
"""Export Markdown draft to Word (.docx) using python-docx."""

import re
import sys
from pathlib import Path

import yaml

try:
    from docx import Document
    from docx.shared import Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Markdown → structured sections
# ---------------------------------------------------------------------------

def _parse_md(text: str) -> list:
    """Return list of (level, heading, body_text). level 0 = title."""
    sections, cur_heading, cur_level, cur_lines = [], None, 0, []
    for line in text.splitlines():
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            if cur_heading is not None:
                sections.append((cur_level, cur_heading, "\n".join(cur_lines).strip()))
            cur_level = len(m.group(1))
            cur_heading = m.group(2)
            cur_lines = []
        else:
            if cur_heading is not None:
                cur_lines.append(line)
    if cur_heading is not None:
        sections.append((cur_level, cur_heading, "\n".join(cur_lines).strip()))
    return sections


# ---------------------------------------------------------------------------
# Font helper
# ---------------------------------------------------------------------------

def _set_run_font(run, jp: str, en: str, size: float):
    run.font.name = en
    run.font.size = Pt(size)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), jp)


# ---------------------------------------------------------------------------
# Main export
# ---------------------------------------------------------------------------

def export_to_docx(draft_md: str, cfg: dict, output_path: Path,
                   template_path: Path = None) -> Path:
    ex = cfg.get("export", {})
    jp = ex.get("font_family_jp", "游明朝")
    en = ex.get("font_family_en", "Times New Roman")
    sz_body = float(ex.get("font_size_body", 10.5))
    sz_h = float(ex.get("font_size_heading", 12))
    spacing = float(ex.get("line_spacing", 1.5))
    margin = int(ex.get("margin_mm", 25))

    if template_path and template_path.exists():
        doc = Document(str(template_path))
        for el in list(doc.element.body):
            doc.element.body.remove(el)
    else:
        doc = Document()
        for sec in doc.sections:
            sec.top_margin = Mm(margin)
            sec.bottom_margin = Mm(margin)
            sec.left_margin = Mm(margin)
            sec.right_margin = Mm(margin)

    for level, heading, body in _parse_md(draft_md):
        # Heading
        if level == 0 or level == 1:
            p = doc.add_heading(heading, level=min(level, 1))
            if level == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            p = doc.add_heading(heading, level=2)
        else:
            p = doc.add_heading(heading, level=3)
        for run in p.runs:
            _set_run_font(run, jp, en, sz_h if level <= 2 else sz_body + 0.5)

        # Body paragraphs
        for block in body.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if re.match(r"^[-*] ", block):
                for item in block.splitlines():
                    item_text = re.sub(r"^[-*]\s+", "", item).strip()
                    if item_text:
                        bp = doc.add_paragraph(style="List Bullet")
                        run = bp.add_run(item_text)
                        _set_run_font(run, jp, en, sz_body)
            else:
                pp = doc.add_paragraph()
                pp.paragraph_format.line_spacing = spacing
                run = pp.add_run(block)
                _set_run_font(run, jp, en, sz_body)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: docx_export.py <draft.md> <output.docx> [config.yaml] [template.docx]")
        sys.exit(1)
    md = Path(sys.argv[1]).read_text(encoding="utf-8")
    out = Path(sys.argv[2])
    cfg = yaml.safe_load(Path(sys.argv[3]).read_text(encoding="utf-8")) if len(sys.argv) > 3 else {}
    tmpl = Path(sys.argv[4]) if len(sys.argv) > 4 else None
    result = export_to_docx(md, cfg, out, tmpl)
    print(f"Exported → {result}")
